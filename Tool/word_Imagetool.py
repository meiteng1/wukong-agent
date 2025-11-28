import os
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from langchain.tools import tool
try:
    from PIL import Image
except Exception:
    Image = None


class ImageInserter:
    """
    根据表格中出现的图片字段（如 xxx.jpg），自动从 static 目录读取实际图片并插入。
    要求：
        1. 不改变模板格式与表格排版；
        2. 匹配成功打印日志；
        3. 匹配不到不替换；
        4. 图片宽度统一 5cm；
    """

    def __init__(self, static_dir: str = "static"):
        self.static_dir = static_dir
        if not os.path.exists(self.static_dir):
            print(f"[警告] 静态资源目录不存在: {self.static_dir}")

    def _is_image_field(self, text: str) -> bool:
        """
        判断单元格的内容是否包含图片字段，例如： xxx.jpg, xxx.png
        """
        if not text:
            return False
        text_lower = text.lower()
        return any(text_lower.endswith(ext) for ext in [".jpg", ".jpeg", ".png"])

    def _find_image_path(self, filename: str):
        """
        在 static 目录中查找文件名匹配的图片。
        """
        for root, _, files in os.walk(self.static_dir):
            for file in files:
                if file.lower() == filename.lower():
                    full_path = os.path.join(root, file)
                    return full_path
        return None

    def _convert_image_safe(self, img_path: str, out_dir: str) -> str:
        if Image is None:
            return img_path
        try:
            os.makedirs(out_dir, exist_ok=True)
            base = os.path.basename(img_path)
            name, ext = os.path.splitext(base)
            safe_path = os.path.join(out_dir, f"{name}.jpg")
            with Image.open(img_path) as im:
                im = im.convert('RGB')
                im.save(safe_path, format='JPEG')
            return safe_path
        except Exception:
            return img_path

    def replace_image_fields(self, docx_path: str, output_path: str):
        """
        主功能：读取 docx，查找表格中的图片字段并替换为实际图片。
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"模板文件不存在: {docx_path}")

        print(f"[INFO] 正在处理文档：{docx_path}")

        doc = Document(docx_path)

        # 修复中文字体（保持模板一致）
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 遍历所有表格和单元格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    if not self._is_image_field(cell_text):
                        continue  # 非图片字段，跳过

                    print(f"[匹配尝试] 表格{table_idx+1} 第{row_idx+1}行 第{cell_idx+1}列 字段内容: {cell_text}")

                    img_path = self._find_image_path(cell_text)

                    if img_path:
                        print(f"[匹配成功] 找到图片：{img_path}")

                        # 清空单元格
                        cell.text = ""

                        # 插入图片（宽度统一为 5cm），不合法图片进行安全转换
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        try:
                            run.add_picture(img_path, width=Cm(5))
                        except Exception:
                            safe_img = self._convert_image_safe(img_path, os.path.join(self.static_dir, "_converted"))
                            run.add_picture(safe_img, width=Cm(5))

                        print(f"[替换成功] 已将图片插入表格{table_idx+1} ({row_idx+1},{cell_idx+1})")
                    else:
                        print(f"[未找到匹配图片] 字段 {cell_text} 在 static 目录中无对应文件")

        # 确保输出目录存在
        out_dir = os.path.dirname(os.path.abspath(output_path)) or os.getcwd()
        os.makedirs(out_dir, exist_ok=True)
        # 保存输出文件（带权限回退）
        try:
            doc.save(output_path)
            print(f"[完成] 已输出到：{output_path}")
            return output_path
        except PermissionError:
            base, ext = os.path.splitext(os.path.abspath(output_path))
            from datetime import datetime
            alt = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
            doc.save(alt)
            print(f"[完成] 已输出到：{alt}")
            return alt


# ========= 将工具封装为 LangChain Tool========= #

@tool
def insert_images_to_docx(template_path: str, output_path: str, static_dir: str = "static") -> str:
    """
    根据模板中的图片字段（xxx.jpg / xxx.png），自动从 static 目录匹配插入图片，并保持模板格式。

    参数:
        template_path: 模板报告路径
        output_path: 输出文件路径
        static_dir: 静态资源目录

    返回：
        输出文件路径
    """
    inserter = ImageInserter(static_dir=static_dir)
    return inserter.replace_image_fields(template_path, output_path)
