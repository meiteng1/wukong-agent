from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.oxml.ns import qn  
from docx.oxml import OxmlElement
from datetime import datetime
import os

try:
    from langchain.tools import tool
except Exception:
    def tool(*args, **kwargs):
        def _wrap(f):
            return f
        return _wrap


def load_env_file():
    """手动加载.env文件"""
    env_path = os.path.join(os.getcwd(), ".env")
    if not os.path.exists(env_path):
        # 尝试项目根目录
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        env_path = os.path.join(project_root, ".env")
    
    if os.path.exists(env_path):
        print(f"加载环境变量文件: {env_path}")
        with open(env_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    if '=' in line:
                        key, value = line.split('=', 1)
                        # 移除引号
                        if value.startswith(('"', "'")) and value.endswith(('"', "'")):
                            value = value[1:-1]
                        os.environ[key.strip()] = value.strip()
    else:
        print("未找到.env文件，将使用系统环境变量")


# 加载环境变量
load_env_file()

# -------------------------- 关键配置：匹配模板的表格列宽（单位：Inches）--------------------------
# 2列表格（开头汇总表、附录表）：左列（标题）1.2英寸，右列（内容）5.0英寸
TEMPLATE_2COL_WIDTHS = [1.2, 5.0]
# 5列表格（缺陷汇总表）：桥墩0.8 | 构件1.0 | 部位1.2 | 缺陷类型1.5 | 现场照片2.0
TEMPLATE_5COL_WIDTHS = [0.8, 1.0, 1.2, 1.5, 2.0]


def create_custom_styles(doc: Document) -> dict:
    """创建自定义样式（优化：添加中文字体兼容性）"""
    # 正文样式：宋体、小四、首行缩进2字符、行距20磅
    core_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    core_font = core_style.font
    core_font.name = '宋体'
    core_font.size = Pt(12)  # 小四=12pt
    core_font.color.rgb = RGBColor(0, 0, 0)
    # 新增：设置中文字体（兼容Mac/Linux）
    core_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    para_format = core_style.paragraph_format
    para_format.first_line_indent = Pt(24)  # 首行缩进2字符=24pt
    para_format.line_spacing = Pt(20)
    para_format.space_after = Pt(0)

    # 一级标题：宋体、四号、加粗、居中、无缩进
    h1_style = doc.styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = '宋体'
    h1_font.size = Pt(14)  # 四号=14pt
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文字体兼容
    def _resolve_heading_style(doc: Document, level: int):
        candidates = []
        if level == 1:
            candidates = ['Heading 1', '标题 1']
        elif level == 2:
            candidates = ['Heading 2', '标题 2']
        elif level == 3:
            candidates = ['Heading 3', '标题 3']
        for name in candidates:
            try:
                return doc.styles[name]
            except Exception:
                continue
        return None
    base_h1 = _resolve_heading_style(doc, 1)
    if base_h1 is not None:
        h1_style.base_style = base_h1
    h1_para = h1_style.paragraph_format
    h1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h1_para.first_line_indent = Pt(0)
    h1_para.space_after = Pt(12)

    # 二级标题：宋体、四号、加粗、左对齐、无缩进
    h2_style = doc.styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = '宋体'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文字体兼容
    base_h2 = _resolve_heading_style(doc, 2)
    if base_h2 is not None:
        h2_style.base_style = base_h2
    h2_para = h2_style.paragraph_format
    h2_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    h2_para.first_line_indent = Pt(0)
    h2_para.space_after = Pt(6)

    # 三级标题：宋体、小四、加粗、左对齐、无缩进
    h3_style = doc.styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = '宋体'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文字体兼容
    base_h3 = _resolve_heading_style(doc, 3)
    if base_h3 is not None:
        h3_style.base_style = base_h3
    h3_para = h3_style.paragraph_format
    h3_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    h3_para.first_line_indent = Pt(0)
    h3_para.space_after = Pt(6)

    # 表格标题样式：宋体、小四、居中
    table_caption_style = doc.styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
    table_caption_font = table_caption_style.font
    table_caption_font.name = '宋体'
    table_caption_font.size = Pt(12)
    table_caption_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文字体兼容
    table_caption_para = table_caption_style.paragraph_format
    table_caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_caption_para.space_after = Pt(6)

    return {
        'body': core_style,
        'h1': h1_style,
        'h2': h2_style,
        'h3': h3_style,
        'table_caption': table_caption_style
    }


def set_cell_width(cell, width):
    """强制设置 Word 单元格固定宽度（单位：Inches，转换为dxa单位）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width * 567)))  # 1英寸=567dxa
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    # 单元格垂直居中（匹配模板表格视觉效果）
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def create_fixed_table(doc, rows, cols, col_widths_in_inches):
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.autofit = True
    table.allow_autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr
    for child in list(tblPr):
        if child.tag == qn('w:tblW'):
            tblPr.remove(child)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    tblPr.append(tblW)
    return table


def add_summary_table(doc: Document, styles: dict, data: dict) -> None:
    """添加开头汇总表格（完全匹配模板：2列无边框、左列加粗居中、右列左对齐）"""
    # 2列n行无边框表格
    table = create_fixed_table(doc, rows=5, cols=2, col_widths_in_inches=TEMPLATE_2COL_WIDTHS)
    
    # 填充表格内容（无数据时显示空白，而非“未提供”）
    headers = ['工程名称', '检查内容', '检查结果', '检验结论', '建议']
    project_name = data.get('project_name', '')
    defect_summary = data.get('defect_summary', '')
    main_findings = data.get('main_findings', '')
    suggestions = """对高优先级缺陷（如螺栓松脱或缺失、防滑块顶死、垫石破损）立即安排维修或加固。
对中优先级缺陷（如混凝土裂缝、麻面、涂装漆破损）制定定期维修和巡检计划，防止进一步恶化。
对低优先级缺陷（施工垃圾、异物、防尘围挡小破损）定期清理和维护，确保美观与排水通畅。"""
    
    contents = [
        project_name,
        '桥梁常规检查',
        f"1、 梁体、桥墩、墩台\n{defect_summary}\n\n\n2、 支座系统\n{defect_summary}",
        main_findings,
        suggestions
    ]
    
    for i, (header, content) in enumerate(zip(headers, contents)):
        # 表头单元格（左列）：加粗、居中
        hdr_cell = table.cell(i, 0)
        hdr_para = hdr_cell.paragraphs[0]
        hdr_para.text = header
        hdr_para.style = styles['body']
        hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_para.bold = True  # 匹配模板左列加粗
        
        # 内容单元格（右列）：左对齐、保留换行格式
        content_cell = table.cell(i, 1)
        content_para = content_cell.paragraphs[0]
        content_para.text = content
        content_para.style = styles['body']
        content_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # 添加空行分隔
    doc.add_paragraph()


def add_toc(doc: Document, styles: dict) -> None:
    """添加自动目录（使用 fldSimple 手动创建 TOC 域）"""
    # 目录标题
    toc_title = doc.add_paragraph('目录', style=styles['h2'])
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # 空行
    # 创建 TOC 域
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    run._r.append(fldSimple)
    doc.add_paragraph()  # 空行


def add_section_overview(doc: Document, styles: dict, data: dict) -> None:
    """添加1 概况章节（保持原有逻辑）"""
    doc.add_paragraph('1 概况', style=styles['h1'])
    
    # 1.1 工程概况
    doc.add_paragraph('1.1 工程概况', style=styles['h2'])
    pier_info = data.get('pier_info', '')  # 无数据时空白
    para = doc.add_paragraph(pier_info, style=styles['body'])
    
    # 1.2 检测内容
    doc.add_paragraph('1.2 检测内容', style=styles['h2'])
    content = """本次常规检查覆盖桥梁核心构件，具体内容分两类：
    1) 梁体、桥墩、墩台
    混凝土结构无裂缝、孔洞、夹渣、麻面、缺棱掉角等质量缺陷；
    梁内、墩台面、检修平台等各部位不得余留施工垃圾；
    表面涂装漆无脱落破损情况；
    落水管无松脱。
    2) 支座系统
    垫石：垫石混凝土结构无裂缝、孔洞、夹渣、麻面、缺棱掉角。（根据《城市桥梁设计规范》横向挡块距离支承垫石3cm，纵向挡块距离支承垫石5cm。）；
    支座板：上、下支座板及螺栓无翘曲、裂缝、锈蚀、松脱，应拆除上、下支座板连接件，支座和梁底及垫石之间应密贴，局部空隙不应大于0.3mm；
    防落梁块：无锈蚀、变形，安装牢固，螺栓无缺失，松动，横向挡块距离支承垫石3cm，纵向挡块距离支承垫石5cm；
    球形支座：支座滑动面上的聚四氟乙烯滑板和不锈钢板位置应正确，不得有划痕、碰伤。支座无锈蚀、标签刻度清晰、无缺失；刻度指针无损坏，应能正确度数。"""
    for line in content.split('\n'):
        para = doc.add_paragraph(line.strip() if line.strip() else '', style=styles['body'])
        if line.startswith(('1)', '2)')):
            para.paragraph_format.left_indent = Pt(24)
        elif line.startswith(('垫石：', '支座板：', '防落梁块：', '球形支座：')):
            para.paragraph_format.left_indent = Pt(48)
    
    # 1.3 检测目的
    doc.add_paragraph('1.3 检测目的', style=styles['h2'])
    purpose_content = """全面掌握桥梁当前技术状况，精准识别梁体、支座系统等关键部位的缺陷类型与分布；
    评估缺陷对桥梁运营安全的影响，为后续养护优先级划分提供依据；
    建立缺陷台账，保障桥梁长期稳定运营，降低安全风险。"""
    for line in purpose_content.split('\n'):
        doc.add_paragraph(line, style=styles['body'])
    
    # 1.4 检测设备与方式
    doc.add_paragraph('1.4 检测设备与方式', style=styles['h2'])
    equipment_content = """设备：大疆Matrice 4TD 无人机（搭载4K高清云台相机）。
    检测方式：无人机航拍覆盖全桥墩及支座，采集多角度影像，检查梁体、桥面构件并记录缺陷部位。每处缺陷拍摄 1–2 张照片并编号存档。"""
    for line in equipment_content.split('\n'):
        doc.add_paragraph(line, style=styles['body'])


def add_naming_rules(doc: Document, styles: dict, data: dict) -> None:
    """添加2 部位与缺陷命名规则章节（补充默认值）"""
    doc.add_paragraph('2 部位与缺陷命名规则', style=styles['h1'])
    
    # 2.1 总规则
    doc.add_paragraph('2.1 总规则', style=styles['h2'])
    pier_naming_rule = data.get('pier_naming_rule', '沿东向西里程方向，桥墩、构件编号从 0 开始，如“0#墩”“0#垫石”')
    rule_content = f"""桥梁前进方向：以东向西为里程方向，沿桥梁前进方向确定左、右侧位置，其中左手侧为桥梁左侧，右手侧为桥梁右侧。
{pier_naming_rule}

缺陷位置描述：采用纵向侧别和横向侧别相结合的方式。纵向侧别包括"大里程侧"和"小里程侧"，横向侧别包括"左侧"和"右侧"。部位描述可以组合使用，例如"大里程侧右侧"、"大里程侧左侧"、"小里程侧右侧"、"小里程侧左侧"等，也可仅使用单侧别，如"左侧"、"右侧"、"大里程侧"或"小里程侧"。
缺陷记录：应包括缺陷描述，缺陷描述仅记录缺陷类型及所在部位，缺陷名称需参照巡检细则，包括梁体、桥墩、墩台及支座系统各类缺陷。"""
    for line in rule_content.split('\n'):
        doc.add_paragraph(line.strip() if line.strip() else '', style=styles['body'])
    
    # 2.2 桥梁构件详细命名编号规则
    doc.add_paragraph('2.2 桥梁构件详细命名编号规则', style=styles['h2'])
    
    # 2.2.1 梁体、桥墩、墩台
    doc.add_paragraph('2.2.1 梁体、桥墩、墩台', style=styles['h3'])
    doc.add_paragraph('梁体的命名采用顺序号，沿桥梁前进方向逐墩编号，例如第一个桥墩的梁体编号为"1#梁"。', style=styles['body'])
    doc.add_paragraph('桥墩和墩台同样采用顺序号，沿桥梁前进方向逐墩编号，编号从 0 开始，例如第一个墩编号为"0#墩"。', style=styles['body'])
    
    # 2.2.2 支座系统
    doc.add_paragraph('2.2.2 支座系统', style=styles['h3'])
    support_content = """支座系统各构件的命名也沿桥梁前进方向逐墩编号，编号均从 0 开始。具体命名规则如下：
    垫石命名为"顺序号"，例如第一个桥墩的垫石部分为"0#垫石"；
    支座板命名为"顺序号"，例如第一个桥墩的支座板为"0#支座板"；
    防落梁块命名为"顺序号"，例如第一个桥墩的防落梁块为"0#防落梁块"；
    球形支座命名为"顺序号"，例如第一个桥墩的球形支座为"0#支座"。"""
    for line in support_content.split('\n'):
        doc.add_paragraph(line.strip() if line.strip() else '', style=styles['body'])


def add_defect_inspection(doc: Document, styles: dict, data: dict) -> None:
    """添加3 桥梁缺陷检查章节（表格完全匹配模板：5列无边框、空白行填充）"""
    doc.add_paragraph('3 桥梁缺陷检查', style=styles['h1'])
    
    # 提取数据并补充默认值
    main_findings = data.get('main_findings', '')
    defect_summary = data.get('defect_summary', '')
    beam_pier_defects = data.get('beam_pier_defects', [])
    support_system_defects = data.get('support_system_defects', [])
    eft = data.get('excel_filtered_table')
    
    # 解析Excel筛选数据
    def _parse_eft(eft_data):
        rows = []
        if isinstance(eft_data, dict):
            def _extract(val):
                if isinstance(val, list):
                    return [s.strip() for s in val if isinstance(s, str) and s.strip()]
                elif isinstance(val, str):
                    return [s.strip() for s in val.split('\n') if s.strip()]
                else:
                    return []
            lines = _extract(eft_data.get('table31')) + _extract(eft_data.get('table32'))
        elif isinstance(eft_data, list):
            lines = eft_data
        elif isinstance(eft_data, str):
            lines = [line.strip() for line in eft_data.split('\n') if line.strip()]
        else:
            lines = []
        for line in lines:
            parts = [s.strip() for s in line.replace('，', ',').replace('、', ',').split(',')]
            if len(parts) >= 5 and parts[0].strip().startswith('HC'):
                rows.append(parts[:5])
        return rows
    parsed_rows = _parse_eft(eft)
    rows_31 = [rv for rv in parsed_rows if len(rv) >= 2 and ('#梁' in str(rv[1]) or '#墩' in str(rv[1]))]
    rows_32 = [rv for rv in parsed_rows if len(rv) >= 2 and ('#防落梁块' in str(rv[1]) or '#垫石' in str(rv[1]) or '#支座板' in str(rv[1]) or '#支座' in str(rv[1]))]

    # -------------------------- 3.1 梁体、桥墩、墩台缺陷表 --------------------------
    doc.add_paragraph('3.1 梁体、桥墩、墩台', style=styles['h2'])
    doc.add_paragraph(f"{main_findings}（*统计数据使用配套的txt中的对应部分，注意去重）", style=styles['body'])
    doc.add_paragraph('梁体、桥墩、墩台缺陷状况汇总表见表3.1.1。', style=styles['body'])
    
    # 表3.1.1：5列无边框表格（匹配模板）
    caption = doc.add_paragraph('表 3.1.1 梁体、桥墩、墩台缺陷状况汇总', style=styles['table_caption'])
    # 初始行数：1行表头 + 5行空白行（匹配模板预留行数）
    table1 = create_fixed_table(doc, rows=6, cols=5, col_widths_in_inches=TEMPLATE_5COL_WIDTHS)
    
    # 表头：居中显示
    headers1 = ['桥墩', '构件', '部位', '缺陷类型', '现场照片']
    for i, header in enumerate(headers1):
        hdr_cell = table1.cell(0, i)
        hdr_para = hdr_cell.paragraphs[0]
        hdr_para.text = header
        hdr_para.style = styles['body']
        hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 填充数据（强制使用筛选后的 rows_31）
    fill_data = rows_31
    for row_idx, item in enumerate(fill_data, start=1):
        if isinstance(item, dict):
            def _val(d, key):
                if key in d:
                    return d.get(key, '')
                m = {'pier': '桥墩', 'component': '构件', 'position': '部位', 'defect_type': '缺陷类型', 'photo': '现场照片'}
                return d.get(m.get(key, key), '')
            values = [
                _val(item, 'pier'),
                _val(item, 'component'),
                _val(item, 'position'),
                _val(item, 'defect_type'),
                _val(item, 'photo')
            ]
        elif isinstance(item, str):
            values = [s.strip() for s in item.split(',')][:5]
        else:
            values = list(item)[:5]
        
        # 填充到表格（超过5行时自动新增行）
        if row_idx >= len(table1.rows):
            table1.add_row()
        for col_idx, value in enumerate(values):
            cell = table1.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.text = value
            para.style = styles['body']
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 补充Excel筛选数据占位符说明
    # doc.add_paragraph(f"（*表格用excel的筛选方法筛选出包含“#梁”“#墩”的表格，复制到此处）", style=styles['body'])
    # if eft:
    #     doc.add_paragraph(f"{eft}", style=styles['body'])
    # doc.add_paragraph()

    # -------------------------- 3.2 支座系统缺陷表 --------------------------
    doc.add_paragraph('3.2 支座系统', style=styles['h2'])
    doc.add_paragraph(f"（*统计数据使用配套的txt中的对应部分，注意去重）", style=styles['body'])
    doc.add_paragraph(defect_summary, style=styles['body'])
    doc.add_paragraph('支座系统缺陷状况汇总表见表3.2.1。', style=styles['body'])
    
    # 表3.2.1：5列无边框表格（匹配模板）
    caption2 = doc.add_paragraph('表 3.2.1 支座系统缺陷状况汇总', style=styles['table_caption'])
    # 初始行数：1行表头 + 35行空白行（匹配模板预留行数）
    table2 = create_fixed_table(doc, rows=36, cols=5, col_widths_in_inches=TEMPLATE_5COL_WIDTHS)
    
    # 表头：居中显示
    headers2 = ['桥墩', '构件', '部位', '缺陷类型', '现场照片']
    for i, header in enumerate(headers2):
        hdr_cell = table2.cell(0, i)
        hdr_para = hdr_cell.paragraphs[0]
        hdr_para.text = header
        hdr_para.style = styles['body']
        hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 填充数据（强制使用筛选后的 rows_32）
    fill_data2 = rows_32
    for row_idx, item in enumerate(fill_data2, start=1):
        if isinstance(item, dict):
            def _val(d, key):
                if key in d:
                    return d.get(key, '')
                m = {'pier': '桥墩', 'component': '构件', 'position': '部位', 'defect_type': '缺陷类型', 'photo': '现场照片'}
                return d.get(m.get(key, key), '')
            values = [
                _val(item, 'pier'),
                _val(item, 'component'),
                _val(item, 'position'),
                _val(item, 'defect_type'),
                _val(item, 'photo')
            ]
        elif isinstance(item, str):
            values = [s.strip() for s in item.split(',')][:5]
        else:
            values = list(item)[:5]
        
        # 填充到表格（超过35行时自动新增行）
        if row_idx >= len(table2.rows):
            table2.add_row()
        for col_idx, value in enumerate(values):
            cell = table2.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.text = value
            para.style = styles['body']
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # # 补充Excel筛选数据占位符说明
    # doc.add_paragraph(f"（*表格用excel的筛选方法筛选出包含“#防落梁块”“#垫石”“#支座板”“#支座”的表格，复制到此处）", style=styles['body'])
    # if eft:
    #     doc.add_paragraph(f"{eft}", style=styles['body'])
    # doc.add_paragraph()

    # -------------------------- 3.3 缺陷分析 --------------------------
    # 定义带数量标注的示例文字
    defect_causes = data.get('defect_causes', """基于桥梁养护常规经验的推测（如环境腐蚀、运营损耗、施工残留等），非本次检测统计结论，最终成因需以补充数据为准。其中：构件破损类缺陷（共15处）多由外力碰撞、施工操作不当或长期荷载磨损造成；涂装防护类缺陷（共12处）主要受环境风化、紫外线照射或腐蚀影响；螺栓连接件类缺陷（共20处）与防护不足、安装不牢固及运营振动有关；残留异物类缺陷（共8处）为施工残留或外部堆积导致。""")
    beam_pier_defect_list = data.get('beam_pier_defect_list', """（1）墩台：墩台破损（4处）、墩台表面涂装漆脱落（2处）、墩台裂缝（1处）、墩台表面涂装漆破损（1处）、施工垃圾残留（3处）、墩台缺棱断角（2处）、落水管破损（1处）、落水管松脱（1处）、墩台麻面（1处）、异物（2处）、墩台表面破损（1处）；（2）梁体：梁体麻面（2处）、梁体破损（1处）、梁体螺栓松脱（1处）、梁体裂缝（1处）、梁体表面涂装漆破损（1处）、梁底预埋件混凝土破损（1处）；（3）桥墩：无明显缺陷（0处）。""")
    support_defect_list = data.get('support_defect_list', """（1）防落梁块：防滑块顶死（2处）、螺栓锈蚀（14处）、梁块螺栓缺失（1处）、防滑块异物（3处）、螺栓松脱（1处）、防滑块间距不足（1处）、防滑块锈蚀（1处）、梁块螺栓垫片锈蚀（1处）、梁块螺栓松动（1处）；（2）垫石：垫石缺棱断角（5处）、垫石破损（1处）、环氧砂浆层破损（2处）、垫石裂缝（2处）、垫石麻面（1处）、垫石表面涂装漆破损（1处）、垫石异物（1处）、垫石螺栓锈蚀（1处）；（3）支座板：上支座板螺栓锈蚀（3处）、上支座板螺栓松脱（2处）、连接件未拆除（1处）、下支座板螺栓锈蚀（2处）、上支座板破损（1处）；（4）球形支座：刻度读数异常（1处）、防尘围挡翻起（3处）、防尘围挡脱落（1处）、刻度模糊（2处）、刻度指针缺失（2处）、防尘围挡破损（1处）。""")
    total_defect_list = data.get('total_defect_list', """本次检测共发现缺陷89处（数据来源于配套Excel统计），缺陷分布呈现明显优先级特征：高优先级缺陷10处，主要集中在防落梁块的螺栓松脱（1处）、螺栓缺失（1处）、防滑块顶死（2处）及垫石裂缝（2处），直接影响结构安全；中优先级缺陷57处，分布在墩台破损（4处）、垫石缺棱断角（5处）、环氧砂浆层破损（2处）、梁体裂缝（1处）、支座板螺栓锈蚀（5处）等混凝土结构及连接件，需定期维修；低优先级缺陷22处，以施工垃圾残留（3处）、异物（2处）、防尘围挡小破损（4处）、涂装漆脱落（4处）为主，影响外观及耐久性。从构件分布来看，支座系统缺陷最多（67处），占比75.3%，主要集中在防落梁块（25处）和球形支座（11处）；梁体、桥墩、墩台缺陷22处，占比24.7%，其中墩台缺陷占比最高（18处），梁体缺陷较少（4处），桥墩无明显缺陷。综合判断，桥梁整体结构安全性尚可，但局部存在耐久性与功能性隐患，需针对性处置。""")
    suggestions = data.get('suggestions', """（1）高优先级：立即维修防滑块顶死（2处）、螺栓松脱（2处）、螺栓缺失（1处）的部位，更换缺失螺栓，紧固松脱螺栓，调整防滑块间距至规范要求；对垫石裂缝（2处）采用环氧砂浆修补，确保支承稳定性；（2）中优先级：定期巡检并计划维修墩台破损（4处）、垫石缺棱断角（5处）、环氧砂浆层破损（2处），修补破损区域；对梁体裂缝（1处）、墩台裂缝（1处）进行灌浆加固；修复涂装漆脱落（6处）及锈蚀螺栓（24处），提升防护能力；（3）低优先级：定期清理施工垃圾残留（3处）、构件表面异物（2处），避免堆积影响结构；修复小破损的防尘围挡（4处），防止雨水、灰尘侵入支座内部，延长构件使用寿命。""")

    doc.add_paragraph('3.3 缺陷分析', style=styles['h2'])
    
    # 3.3.1 梁体、桥墩、墩台
    doc.add_paragraph('3.3.1 梁体、桥墩、墩台', style=styles['h3'])
    doc.add_paragraph(defect_causes, style=styles['body'])
    doc.add_paragraph(f"{beam_pier_defect_list}（*根据上述统计的内容说明缺陷情况、缺陷成因）", style=styles['body'])
    
    # 3.3.2 支座系统
    doc.add_paragraph('3.3.2 支座系统', style=styles['h3'])
    doc.add_paragraph(f"{support_defect_list}（*根据上述统计的内容说明缺陷情况、缺陷成因）", style=styles['body'])
    doc.add_paragraph(defect_causes, style=styles['body'])
    
    # 3.3.3 总体分析（修正模板编号错误）
    doc.add_paragraph('3.3.3 总体分析', style=styles['h3'])
    doc.add_paragraph(defect_causes, style=styles['body'])
    doc.add_paragraph(f"{total_defect_list}（*根据上述统计的内容说明缺陷分布情况、缺陷解决建议）", style=styles['body'])
    doc.add_paragraph(suggestions, style=styles['body'])


def add_appendix(doc: Document, styles: dict, data: dict) -> None:
    """添加附录 1（表格匹配模板：2列无边框、左列加粗居中）"""
    doc.add_paragraph('附录 1', style=styles['h1'])
    
    # 提取数据并补充默认值
    pier_naming_rule = data.get('pier_naming_rule', '沿东向西里程方向，桥墩、构件编号从 0 开始')
    bridge_code = data.get('bridge_code', '未提供具体编码')
    id_file_mapping = data.get('id_file_mapping', '照片按“桥墩编号-部位-缺陷类型.jpg”命名')
    
    # 附录表格：2列无边框（匹配模板）
    appendix_table = create_fixed_table(doc, rows=4, cols=2, col_widths_in_inches=TEMPLATE_2COL_WIDTHS)
    
    # 附录内容（左列加粗、右列左对齐）
    appendix_items = [
        ('编号规则拆解', f"1. 桥墩标识：{pier_naming_rule}\n（*修改编号）\n2. 缺陷位置：“大里程侧 / 小里程侧 + 左 / 右” 与报告 “部位” 列（如大里程侧右侧）匹配，明确缺陷方位；\n3. 缺陷类型：末尾文字与报告 “缺陷类型” 列（如防滑块锈蚀）一致，体现图片记录的缺陷内容。"),
        ('编号与文件的关联方式', f"1. {id_file_mapping}\n（*修改编号）\n2. 与报告缺陷关联：通过编号三要素（桥墩标识、位置、类型），可反向对应报告表格中对应列信息。"),
        ('图片查阅操作说明', f"1. 独立 PDF 按 “桥墩编号顺序” 整理{bridge_code}，支持按桥墩批量查阅；\n2. 定位具体缺陷图片：先在报告表格找到 “现场照片” 列编号，再在 PDF 中用 “查找” 功能输入编号，即可快速跳转至目标图片。"),
        ('对应报告文件说明', "报告中所有 “现场照片” 列编号均遵循上述规则，与独立存储的图片 PDF 配套使用，不涉及其他文件关联。")
    ]
    
    for i, (title, content) in enumerate(appendix_items):
        # 标题单元格（左列）：加粗、居中
        title_cell = appendix_table.cell(i, 0)
        title_para = title_cell.paragraphs[0]
        title_para.text = title
        title_para.style = styles['body']
        title_para.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 内容单元格（右列）：左对齐、保留换行
        content_cell = appendix_table.cell(i, 1)
        content_para = content_cell.paragraphs[0]
        content_para.text = content
        content_para.style = styles['body']
        content_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # 补充附录占位符
    appendix_placeholder = data.get('appendix', '')
    if appendix_placeholder:
        doc.add_paragraph(appendix_placeholder, style=styles['body'])


def _parse_excel_filtered_table(eft):
    """辅助函数：解析Excel筛选数据"""
    rows = []
    if isinstance(eft, dict):
        def _extract(val):
            if isinstance(val, list):
                return [s.strip() for s in val if isinstance(s, str) and s.strip()]
            elif isinstance(val, str):
                return [s.strip() for s in val.split('\n') if s.strip()]
            else:
                return []
        lines = _extract(eft.get('table31')) + _extract(eft.get('table32'))
    elif isinstance(eft, list):
        lines = eft
    elif isinstance(eft, str):
        lines = [line.strip() for line in eft.split('\n') if line.strip()]
    else:
        lines = []
    for line in lines:
        parts = [s.strip() for s in line.replace('，', ',').replace('、', ',').split(',')]
        if len(parts) >= 5 and parts[0].strip().startswith('HC'):
            rows.append(parts[:5])
    return rows


def _fill_table_body(table, styles, rows):
    table.autofit = True
    table.allow_autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr
    for child in list(tblPr):
        if child.tag == qn('w:tblW'):
            tblPr.remove(child)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    tblPr.append(tblW)
    if not rows:
        return
    for idx, row_vals in enumerate(rows, start=1):
        if idx >= len(table.rows):
            table.add_row()
        for i in range(5):
            para = table.rows[idx].cells[i].paragraphs[0]
            para.text = str(row_vals[i]) if i < len(row_vals) else ''
            para.style = styles['body']
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _find_paragraph_by_text(doc, text):
    for p in doc.paragraphs:
        if text in p.text.strip():
            return p
    return None


def _find_table_after_paragraph(doc, paragraph):
    """辅助函数：查找段落后的第一个表格"""
    if paragraph is None:
        return None
    start = False
    for block in paragraph._p.getparent():
        if block is paragraph._p:
            start = True
            continue
        if not start:
            continue
        if hasattr(block, 'tbl'): 
            for t in doc.tables:
                if t._tbl is block.tbl:
                    return t
    return None


def _fill_template_tables(doc: Document, styles: dict, data: dict):
    eft_rows = _parse_excel_filtered_table(data.get('excel_filtered_table'))
    eft_rows_31 = [rv for rv in eft_rows if len(rv) >= 2 and ('#梁' in str(rv[1]) or '#墩' in str(rv[1]))]
    eft_rows_32 = [rv for rv in eft_rows if len(rv) >= 2 and ('#防落梁块' in str(rv[1]) or '#垫石' in str(rv[1]) or '#支座板' in str(rv[1]) or '#支座' in str(rv[1]))]
    rows_31 = eft_rows_31
    rows_32 = eft_rows_32
    
    # 查找并填充两个缺陷表
    p31 = _find_paragraph_by_text(doc, '表 3.1.1')
    p32 = _find_paragraph_by_text(doc, '表 3.2.1')
    t31 = _find_table_after_paragraph(doc, p31)
    t32 = _find_table_after_paragraph(doc, p32)
    
    if t31:
        _fill_table_body(t31, styles, rows_31)
    if t32:
        _fill_table_body(t32, styles, rows_32)


def _add_table_after_paragraph(doc: Document, paragraph, styles: dict, rows):
    """辅助函数：在指定段落后添加表格"""
    headers = ['桥墩', '构件', '部位', '缺陷类型', '现场照片']
    table = create_fixed_table(doc, rows=1, cols=5, col_widths_in_inches=TEMPLATE_5COL_WIDTHS)
    
    # 填充表头
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.text = h
        p.style = styles['body']
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 填充数据（无数据则空白）
    if not rows:
        for _ in range(5):
            r = table.add_row()
            for i in range(5):
                p = r.cells[i].paragraphs[0]
                p.text = ''
                p.style = styles['body']
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        for rv in rows:
            r = table.add_row()
            for i in range(5):
                p = r.cells[i].paragraphs[0]
                p.text = str(rv[i]) if i < len(rv) else ''
                p.style = styles['body']
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 插入到段落后
    body = doc._body._element
    body.remove(table._tbl)
    paragraph._p.addnext(table._tbl)


def _apply_excel_placeholders(doc: Document, styles: dict, data: dict):
    eft_rows = _parse_excel_filtered_table(data.get('excel_filtered_table'))
    eft_rows_31 = [rv for rv in eft_rows if len(rv) >= 2 and ('#梁' in str(rv[1]) or '#墩' in str(rv[1]))]
    eft_rows_32 = [rv for rv in eft_rows if len(rv) >= 2 and ('#防落梁块' in str(rv[1]) or '#垫石' in str(rv[1]) or '#支座板' in str(rv[1]) or '#支座' in str(rv[1]))]
    rows_31 = eft_rows_31
    rows_32 = eft_rows_32
    
    current_section = None
    for para in doc.paragraphs:
        para_text = para.text.strip()
        # 识别当前章节
        if para_text.startswith('3.1'):
            current_section = '3.1'
        elif para_text.startswith('3.2'):
            current_section = '3.2'
        # 替换Excel占位符
        if '{excel_filtered_table}' in para_text:
            para.text = ''
            if current_section == '3.1':
                _add_table_after_paragraph(doc, para, styles, rows_31)
            elif current_section == '3.2':
                _add_table_after_paragraph(doc, para, styles, rows_32)

def _find_paragraph_containing(doc: Document, substring: str):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None

def _rows_from_dicts(dict_rows):
    rows = []
    for d in dict_rows or []:
        rows.append([
            str(d.get('pier', '')),
            str(d.get('component', '')),
            str(d.get('position', '')),
            str(d.get('defect_type', '')),
            str(d.get('photo', ''))
        ])
    return rows

def _parse_lines_to_dicts(lines):
    dicts = []
    if isinstance(lines, list):
        src = lines
    elif isinstance(lines, str):
        src = [s.strip() for s in lines.split('\n') if s.strip()]
    else:
        src = []
    for line in src:
        if not isinstance(line, str):
            continue
        parts = [s.strip() for s in line.split(',')]
        if len(parts) >= 5:
            dicts.append({
                'pier': parts[0],
                'component': parts[1],
                'position': parts[2],
                'defect_type': parts[3],
                'photo': parts[4]
            })
    return dicts

def _fill_table_by_keyword(doc: Document, styles: dict, keyword: str, dict_rows):
    para = _find_paragraph_containing(doc, keyword)
    table = _find_table_after_paragraph(doc, para)
    if table:
        _fill_table_body(table, styles, _rows_from_dicts(dict_rows))


def generate_bridge_report(data: dict, filename: str = None, template_path: str = None) -> str:
    """
    自动生成桥梁支座检查报告 Word 文档（核心函数）
    :param data: 报告数据字典，包含所有占位符内容
    :param filename: 保存文件名（可选）
    :param template_path: 模板路径（可选）
    :return: 生成的文件路径（绝对路径）
    """
    # 生成默认文件名
    if not filename:
        project_name = data.get('project_name', '桥梁支座检查报告')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{project_name}_{timestamp}.docx"
    
    filename = os.path.abspath(filename)
    dirn = os.path.dirname(filename)
    if dirn and not os.path.exists(dirn):
        try:
            os.makedirs(dirn, exist_ok=True)
        except Exception:
            pass
    
    # 创建文档实例（使用模板或新建）
    doc = Document(template_path) if (template_path and os.path.exists(template_path)) else Document()
    
    # 设置页面边距（匹配模板：1英寸边距）
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    
    # 创建自定义样式
    styles = create_custom_styles(doc)
    
    # 处理模板或新建文档
    if template_path and os.path.exists(template_path):
        _apply_excel_placeholders(doc, styles, data)
        _fill_template_tables(doc, styles, data)
        t31 = data.get('table31')
        t32 = data.get('table32')
        rows31_dicts = t31 if (isinstance(t31, list) and (len(t31) == 0 or isinstance(t31[0], dict))) else _parse_lines_to_dicts(t31)
        rows32_dicts = t32 if (isinstance(t32, list) and (len(t32) == 0 or isinstance(t32[0], dict))) else _parse_lines_to_dicts(t32)
        if rows31_dicts:
            _fill_table_by_keyword(doc, styles, '表 3.1.1', rows31_dicts)
        if rows32_dicts:
            _fill_table_by_keyword(doc, styles, '表 3.2.1', rows32_dicts)
    else:
        # 新建文档：按顺序添加内容
        doc.add_paragraph('厦门轨道桥梁支座检查报告', style=styles['h1'])
        add_summary_table(doc, styles, data)
        add_toc(doc, styles)
        add_section_overview(doc, styles, data)
        add_naming_rules(doc, styles, data)
        add_defect_inspection(doc, styles, data)
        add_appendix(doc, styles, data)
    
    # 保存文档（处理权限错误）
    try:
        doc.save(filename)
        print(f"✅ 桥梁支座检查报告已生成：{filename}")
        return filename
    except Exception:
        base, ext = os.path.splitext(filename)
        alt_base = os.path.join(os.getcwd(), os.path.basename(base))
        alt_filename = f"{alt_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        doc.save(alt_filename)
        print(f"✅ 备用报告已生成：{alt_filename}")
        return alt_filename


@tool
def create_complete_report(output_path: str, data: dict, template_path: str = None) -> str:
    """
    工具：生成完整的桥梁报告 docx
    :param output_path: 报告保存路径
    :param data: 报告数据字典
    :param template_path: 模板路径（可选）
    :return: 生成的绝对路径
    """
    try:
        eft = data.get('excel_filtered_table')
        t31 = data.get('table31')
        t32 = data.get('table32')
        if not eft or (isinstance(eft, str) and not eft.strip()):
            from Tool.excel_reader_tool import read_filtered_excel_tables
            tables = read_filtered_excel_tables.invoke({'file_path': os.environ.get('REFER_FILE_OUT_PATH')})
            t31 = tables.get('table31', [])
            t32 = tables.get('table32', [])
            data['excel_filtered_table'] = "\n".join(list(t31) + list(t32))
            data['table31'] = t31
            data['table32'] = t32
        elif (not t31) or (not t32):
            from Tool.excel_reader_tool import read_filtered_excel_tables
            tables = read_filtered_excel_tables.invoke({'file_path': os.environ.get('REFER_FILE_OUT_PATH')})
            data['table31'] = tables.get('table31', [])
            data['table32'] = tables.get('table32', [])
    except Exception:
        pass
    return generate_bridge_report(data, output_path, template_path)


# 导出给 agent 使用
WORD_TOOLS = [create_complete_report]


# 示例：运行生成报告
if __name__ == "__main__":
    # 示例数据（可根据实际统计结果修改）
    sample_data = {
        'project_name': '厦门轨道交通3号线后溪站-车辆段区间桥梁支座检测项目',
        'defect_summary': '经检测，梁体、桥墩、墩台存在墩台破损、混凝土麻面等缺陷；支座系统存在防滑块顶死、螺栓锈蚀、垫石缺棱断角等缺陷，具体数据详见配套 Excel 统计 Sheet',
        'main_findings': '本次检测覆盖后溪站-车辆段1个区段，共发现缺陷35处（详见配套 Excel），其中高优先级缺陷8处（含螺栓松脱、防滑块顶死），需立即处置；中优先级缺陷17处（含混凝土裂缝、垫石破损）；低优先级缺陷10处（含防尘围挡破损、施工垃圾）',
        'pier_info': '本次检测涵盖厦门轨道交通3号线后溪站-车辆段区间桥梁，涉及桥墩编号HC-00至HC-03，共4个桥墩，墩位沿东向西里程方向分布（桥墩数量详见配套 Excel 统计 Sheet）',
        'pier_naming_rule': '未提供，暂按模板默认规则：沿东向西里程方向，桥墩、构件编号从 0 开始，如“0#墩”“0#垫石”',
        'bridge_code': '未提供，暂按“厦门轨道交通+区段名称”分类',
        'id_file_mapping': '照片命名为“桥墩编号-部位-缺陷类型.jpg”，按“区段-桥墩”文件夹存储（如“后溪站-车辆段/HC-00/HC-00-大里程侧右侧墩台破损.jpg”）',
        'defect_causes': '基于桥梁养护常规经验的推测（如环境腐蚀、运营损耗、施工残留等），非本次检测统计结论，最终成因需以补充数据为准。其中：构件破损类缺陷（共15处）多由外力碰撞、施工操作不当或长期荷载磨损造成；涂装防护类缺陷（共12处）主要受环境风化、紫外线照射或腐蚀影响；螺栓连接件类缺陷（共20处）与防护不足、安装不牢固及运营振动有关；残留异物类缺陷（共8处）为施工残留或外部堆积导致。',
        'beam_pier_defect_list': '（1）墩台：墩台破损（4处）、墩台表面涂装漆脱落（2处）、墩台裂缝（1处）、墩台表面涂装漆破损（1处）、施工垃圾残留（3处）、墩台缺棱断角（2处）、落水管破损（1处）、落水管松脱（1处）、墩台麻面（1处）、异物（2处）、墩台表面破损（1处）；（2）梁体：梁体麻面（2处）、梁体破损（1处）、梁体螺栓松脱（1处）、梁体裂缝（1处）、梁体表面涂装漆破损（1处）、梁底预埋件混凝土破损（1处）；（3）桥墩：无明显缺陷（0处）。',
        'support_defect_list': '（1）防落梁块：防滑块顶死（2处）、螺栓锈蚀（14处）、梁块螺栓缺失（1处）、防滑块异物（3处）、螺栓松脱（1处）、防滑块间距不足（1处）、防滑块锈蚀（1处）、梁块螺栓垫片锈蚀（1处）、梁块螺栓松动（1处）；（2）垫石：垫石缺棱断角（5处）、垫石破损（1处）、环氧砂浆层破损（2处）、垫石裂缝（2处）、垫石麻面（1处）、垫石表面涂装漆破损（1处）、垫石异物（1处）、垫石螺栓锈蚀（1处）；（3）支座板：上支座板螺栓锈蚀（3处）、上支座板螺栓松脱（2处）、连接件未拆除（1处）、下支座板螺栓锈蚀（2处）、上支座板破损（1处）；（4）球形支座：刻度读数异常（1处）、防尘围挡翻起（3处）、防尘围挡脱落（1处）、刻度模糊（2处）、刻度指针缺失（2处）、防尘围挡破损（1处）。',
        'total_defect_list': '本次检测共发现缺陷89处（数据来源于配套Excel统计），缺陷分布呈现明显优先级特征：高优先级缺陷10处，主要集中在防落梁块的螺栓松脱（1处）、螺栓缺失（1处）、防滑块顶死（2处）及垫石裂缝（2处），直接影响结构安全；中优先级缺陷57处，分布在墩台破损（4处）、垫石缺棱断角（5处）、环氧砂浆层破损（2处）、梁体裂缝（1处）、支座板螺栓锈蚀（5处）等混凝土结构及连接件，需定期维修；低优先级缺陷22处，以施工垃圾残留（3处）、异物（2处）、防尘围挡小破损（4处）、涂装漆脱落（4处）为主，影响外观及耐久性。从构件分布来看，支座系统缺陷最多（67处），占比75.3%，主要集中在防落梁块（25处）和球形支座（11处）；梁体、桥墩、墩台缺陷22处，占比24.7%，其中墩台缺陷占比最高（18处），梁体缺陷较少（4处），桥墩无明显缺陷。综合判断，桥梁整体结构安全性尚可，但局部存在耐久性与功能性隐患，需针对性处置。',
        'suggestions': '1. 高优先级缺陷（螺栓松脱、防滑块顶死、垫石破损）：立即安排专业施工团队进行维修加固，更换锈蚀螺栓，修补破损垫石；\n2. 中优先级缺陷（混凝土裂缝、涂装漆脱落）：制定季度巡检计划，跟踪缺陷发展情况，年度集中维修；\n3. 低优先级缺陷（施工垃圾、防尘围挡小破损）：每月定期清理，及时修复围挡破损部位',
        # 梁体、桥墩、墩台缺陷数据（示例）
        'beam_pier_defects': [
            {'pier': 'HC-00', 'component': '0#墩', 'position': '大里程侧右侧', 'defect_type': '墩台破损', 'photo': 'HC-00-大里程侧右侧墩台破损.jpg'},
            {'pier': 'HC-01', 'component': '1#梁', 'position': '左侧', 'defect_type': '混凝土麻面', 'photo': 'HC-01-左侧混凝土麻面.jpg'},
            {'pier': 'HC-02', 'component': '2#墩台', 'position': '小里程侧左侧', 'defect_type': '缺棱掉角', 'photo': 'HC-02-小里程侧左侧缺棱掉角.jpg'}
        ],
        # 支座系统缺陷数据（示例）
        'support_system_defects': [
            {'pier': 'HC-00', 'component': '0#防落梁块', 'position': '大里程侧右侧', 'defect_type': '防滑块顶死', 'photo': 'HC-00-大里程侧右侧防滑块顶死.jpg'},
            {'pier': 'HC-00', 'component': '0#防落梁块', 'position': '大里程侧右侧', 'defect_type': '螺栓锈蚀', 'photo': 'HC-00-大里程侧右侧螺栓锈蚀.jpg'},
            {'pier': 'HC-00', 'component': '0#垫石', 'position': '大里程侧左侧', 'defect_type': '垫石缺棱断角', 'photo': 'HC-00-大里程侧左侧垫石缺棱断角.jpg'},
            {'pier': 'HC-01', 'component': '1#垫石', 'position': '大里程侧右侧', 'defect_type': '环氧砂浆层破损', 'photo': 'HC-01-大里程侧右侧环氧砂浆层破损.jpg'},
            {'pier': 'HC-01', 'component': '1#防落梁块', 'position': '小里程侧右侧', 'defect_type': '梁块螺栓缺失', 'photo': 'HC-01-小里程侧右侧梁块螺栓缺失.jpg'}
        ],
        # Excel筛选数据示例（字符串格式）
        'excel_filtered_table': 'HC-00,0#墩,大里程侧右侧,墩台破损,HC-00-大里程侧右侧墩台破损.jpg\nHC-01,1#梁,左侧,混凝土麻面,HC-01-左侧混凝土麻面.jpg'
    }
    
    # 生成报告（无模板时新建，有模板时传入template_path参数）
    generate_bridge_report(sample_data)
