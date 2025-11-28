import os
import chardet
from docx import Document
from docx.document import Document as DocObject
from docx.oxml.ns import qn
from langchain.tools import tool

def load_env_file():
    """手动加载.env文件"""
    env_path = os.path.join(os.getcwd(), ".env")
    if not os.path.exists(env_path):
        # 尝试项目根目录
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        env_path = os.path.join(project_root, ".env")
    
    if os.path.exists(env_path):
        print(f"加载环境变量文件: {env_path}")
        with open(env_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    if '=' in line:
                        key, value = line.split('=', 1)
                        # 移除引号
                        if value.startswith(('"', "'")) and value.endswith(('"', "'")):
                            value = value[1:-1]
                        os.environ[key.strip()] = value.strip()
    else:
        print("未找到.env文件，将使用系统环境变量")

# 加载环境变量
load_env_file()

def _parse_docx_tables_to_markdown(table) -> str:
    """
    辅助函数：将docx表格转换为markdown格式（保留原表格结构）
    参数: docx表格对象
    返回: markdown格式的表格字符串
    """
    markdown_table = []
    # 遍历表格的每一行
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        # 遍历行中的每个单元格
        for cell in row.cells:
            # 保留单元格内的换行符（用\n替换），避免内容丢失
            cell_text = cell.text.replace('\r', '').replace('\n', '<br/>')  # 用<br/>保留换行结构
            row_cells.append(cell_text.strip() if cell_text.strip() else "")
        
        # 拼接当前行的markdown格式（|单元格1|单元格2|...|）
        markdown_row = f"|{'|'.join(row_cells)}|"
        markdown_table.append(markdown_row)
        
        # 第一行后添加分隔线（|---|---|...|）
        if row_idx == 0:
            separator = f"|{'|'.join(['---' for _ in row_cells])}|"
            markdown_table.append(separator)
    
    # 表格前后加空行，确保格式清晰，最后加注释保留标记
    newline = '\n'
    return f"{newline}{newline.join(markdown_table)}{newline}"

def _extract_docx_modules(doc: DocObject) -> dict:
    """
    辅助函数：按模板结构提取模块（开头表格、目录、正文章节）
    参数: docx文档对象
    返回: 分模块的内容字典
    """
    modules = {
        "开头表格": "",
        "目录": "",
        "正文章节": ""
    }
    current_module = "开头表格"  # 初始模块：先读取开头表格
    has_extracted_table = False  # 标记是否已提取开头表格
    catalog_start_marker = "目录"  # 目录的起始标记（模板中明确的“目录”标题）
    catalog_end_markers = ["1 概况", "1.1 工程概况"]  # 目录结束的标记（第一章标题）

    # 1. 先提取所有表格（优先处理开头表格）
    for table_idx, table in enumerate(doc.tables):
        table_markdown = _parse_docx_tables_to_markdown(table)
        if not has_extracted_table:
            # 第一个表格即为“开头表格”
            modules["开头表格"] = table_markdown
            has_extracted_table = True
            current_module = "目录"  # 表格提取后，下一个模块是目录
        else:
            # 其他表格（如3.1、3.2的缺陷汇总表）暂归入“正文章节”
            modules["正文章节"] += f"\n【正文表格{table_idx + 1}】\n{table_markdown}"

    # 2. 提取段落内容（按模块分配）
    paragraph_content = []
    for para in doc.paragraphs:
        para_text = para.text.strip()
        para_raw_text = para.text  # 保留原始文本（含空行、缩进）
        # 处理空行：保留原始空行结构（避免丢失模板中的<br/>对应的空行）
        if not para_raw_text:
            paragraph_content.append("")
            continue

        # 模块切换逻辑：从“目录”切换到“正文章节”
        if current_module == "目录":
            # 检测目录结束（遇到第一章标题）
            if any(marker in para_text for marker in catalog_end_markers):
                # 将已收集的目录内容存入“目录”模块
                modules["目录"] = "\n".join(paragraph_content) + f"\n{para_raw_text}"
                paragraph_content = []
                current_module = "正文章节"
            elif para_text == catalog_start_marker:
                # 遇到“目录”标题，开始收集目录内容
                paragraph_content.append(para_raw_text)
            else:
                # 目录内容（如“1 概况1”“1.1 工程概况1”）
                paragraph_content.append(para_raw_text)
        else:
            # 正文章节内容（含所有章节、注释、空行）
            paragraph_content.append(para_raw_text)

    # 3. 处理剩余的段落内容（若目录未完整提取，补充到对应模块）
    if current_module == "目录" and paragraph_content:
        modules["目录"] += "\n".join(paragraph_content)
    elif current_module == "正文章节" and paragraph_content:
        modules["正文章节"] += "\n".join(paragraph_content)

    return modules

def _format_template_preview(modules: dict) -> str:
    """
    辅助函数：将分模块内容格式化为“模板原文预览”字符串
    参数: 分模块的内容字典
    返回: 结构化的预览字符串
    """
    preview_parts = [
        "=" * 50,
        "【模板原文预览-整体说明】",
        "1. 以下内容完整保留模板的表格、空行、注释、章节结构",
        "2. 未标注“修改”的内容均为模板原文，可直接核对完整性",
        "=" * 50
    ]
    # 按模块添加预览内容
    for module_name, content in modules.items():
        if content:
            preview_parts.append(f"\n【模板原文预览-{module_name}】")
            preview_parts.append("-" * 30)
            # 还原表格中的<br/>为换行（方便预览时直观看到换行结构）
            content_formatted = content.replace("<br/>", "\n")
            preview_parts.append(content_formatted)
        else:
            preview_parts.append(f"\n【模板原文预览-{module_name}】")
            preview_parts.append("-" * 30)
            preview_parts.append("未提取到内容（可能模块位置与预期不符，请检查模板结构）")
    
    preview_parts.append("\n" + "=" * 50)
    preview_parts.append("【预览确认提示】请核对上述内容是否与模板完全一致（无表格丢失、无文字遗漏）")
    preview_parts.append("=" * 50)
    return "\n".join(preview_parts)

@tool
def read_text_auto(
    path: str = None,
    is_template_preview: bool = False  # 新增：是否开启模板预览模式
) -> str:
    """
    自动读取txt或docx文件内容（新增模板预览模式，保留表格/空行/注释）
    
    参数:
        path: 文件路径（支持.txt和.docx格式）。如果未提供或文件不存在，将尝试使用环境变量中的TEMPLATE_REPORT_PATH（模板路径）
        is_template_preview: 是否输出“模板原文预览”（结构化分模块，保留完整格式）
    
    返回:
        文件内容字符串（普通模式：纯文本；预览模式：分模块结构化预览）
    """
    # 路径处理：优先使用模板路径（TEMPLATE_REPORT_PATH）
    if not path or not os.path.exists(path):
        # 区分“原始报告”和“模板报告”的环境变量
        env_path = os.getenv("TEMPLATE_REPORT_PATH") or os.getenv("RAW_REPORT_PATH")
        if env_path and os.path.exists(env_path):
            print(f"使用环境变量中的路径: {env_path}")
            path = env_path
        else:
            raise FileNotFoundError(
                f"文件不存在！请检查：1. 输入路径是否正确 2. 环境变量TEMPLATE_REPORT_PATH是否设置（模板路径）"
            )
    
    # 1. 处理docx文件（核心：支持表格读取和模板预览）
    if path.lower().endswith(".docx"):
        doc = Document(path)
        # 修复docx中文乱码问题（设置默认字体）
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        if is_template_preview:
            # 模板预览模式：分模块提取并格式化
            modules = _extract_docx_modules(doc)
            return _format_template_preview(modules)
        else:
            # 普通模式：合并表格和段落（保留基础格式）
            full_content = []
            # 先加表格
            for table_idx, table in enumerate(doc.tables):
                full_content.append(_parse_docx_tables_to_markdown(table))
            # 再加段落
            for para in doc.paragraphs:
                full_content.append(para.text)
            return "\n".join(full_content)
    
    # 2. 处理txt文件（自动检测编码，保留原始格式）
    with open(path, "rb") as file:
        raw_data = file.read()
    
    # 检测文件编码（兼容GBK、UTF-8等）
    detected_encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
    try:
        content = raw_data.decode(detected_encoding)
    except UnicodeDecodeError:
        content = raw_data.decode("utf-8", errors="ignore")
        print(f"警告：使用UTF-8容错模式解码，可能丢失部分特殊字符（原始编码：{detected_encoding}）")
    
    if is_template_preview:
        # 对txt模板也按“开头表格（若有）、目录、正文”分模块预览
        return _format_template_preview({
            "开头表格": "txt文件无表格结构（若有表格请用docx格式）",
            "目录": content.split("1 概况")[0] if "1 概况" in content else "未找到明确目录",
            "正文章节": content.split("1 概况")[1] if "1 概况" in content else content
        })
    return content

@tool
def save_to_docx(content: str, output_path: str) -> str:
    """
    将文本内容保存为docx文件（支持还原markdown表格格式）
    
    参数:
        content: 要保存的文本内容（可含markdown表格、<br/>换行）
        output_path: 输出docx文件路径
    
    返回:
        输出文件路径
    """
    doc = Document()
    # 设置默认字体（避免中文乱码）
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 拆分内容为段落块（按空行拆分，保留大段结构）
    content_blocks = content.split("\n\n")
    for block in content_blocks:
        block = block.strip()
        if not block:
            doc.add_paragraph("")  # 保留空行
            continue
        
        # 检测是否为markdown表格（含"|"和"---"）
        if "|" in block and "---" in block:
            # 解析markdown表格为docx表格
            table_rows = [row.strip() for row in block.split("\n") if row.strip()]
            if len(table_rows) >= 2:  # 至少有表头+分隔线
                # 获取列数（从表头行提取）
                header_cells = [cell.strip() for cell in table_rows[0].strip("|").split("|")]
                col_count = len(header_cells)
                # 创建docx表格
                table = doc.add_table(rows=1, cols=col_count)
                table.style = "Table Grid"  # 带边框的表格样式（还原模板表格外观）
                
                # 填充表头
                header_cells = table.rows[0].cells
                for idx, cell_text in enumerate(header_cells):
                    if idx < len(header_cells):
                        # 还原<br/>为换行符
                        cell_text = cell_text.replace("<br/>", "\n")
                        header_cells[idx].text = cell_text
                
                # 填充表格内容（跳过分隔线行）
                for row_idx, row in enumerate(table_rows[1:]):
                    if "---" in row:
                        continue  # 跳过分隔线
                    row_cells = [cell.strip() for cell in row.strip("|").split("|")]
                    row_cells = row_cells + [""] * (col_count - len(row_cells))  # 补全空单元格
                    table_row = table.add_row().cells
                    for idx, cell_text in enumerate(row_cells):
                        cell_text = cell_text.replace("<br/>", "\n")
                        table_row[idx].text = cell_text
                # 表格后加空行
                doc.add_paragraph("")
                continue
        
        # 普通文本段落（还原<br/>为换行）
        para_text = block.replace("<br/>", "\n")
        para = doc.add_paragraph(para_text)
        # 还原目录/章节的缩进（根据文本前缀判断）
        if para_text.startswith(("1.", "2.", "3.")):  # 一级/二级章节标题
            para.paragraph_format.left_indent = 0  # 无缩进
            para.runs[0].font.bold = True  # 章节标题加粗（还原模板格式）
        elif para_text.startswith(("（*", "注：")):  # 注释文本
            para.paragraph_format.left_indent = 20  # 缩进2字符（突出注释）
    
    doc.save(output_path)
    return f"文件已保存至：{output_path}\n提示：表格已保留边框样式，章节标题已加粗，注释已缩进"

# 保留类形式以便向后兼容（同步更新read_text_auto方法）
class DocumentHandler:
    def __init__(self):
        # 从环境变量获取路径（优先模板路径）
        self.raw_report_path = os.getenv("RAW_REPORT_PATH")
        self.template_report_path = os.getenv("TEMPLATE_REPORT_PATH")  # 新增：模板路径
    
    @staticmethod
    def read_text_auto(path=None, is_template_preview=False):
        return read_text_auto(path=path, is_template_preview=is_template_preview)
    
    @staticmethod
    def save_to_docx(content, output_path):
        return save_to_docx(content, output_path)